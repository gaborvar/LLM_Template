###### 
# This Python script is used to preprocess templates for the LLM based chat service
# which creates documents by combining a template and customer specific data collected interactively in the chat.
# Usage: python templatepreproc.py <path> where <path> points to either a single template or a directory containing templates with docx extension.
# Output is saved in the "PreProc" subdirectory.

from docx import Document
import re, os, sys
import logging

from docx.shared import Pt, RGBColor
from collections import Counter, defaultdict
import unicodedata
import string

logging.basicConfig(level=logging.WARNING)
logger = logging.getLogger("TemplPrep")
logger.setLevel(logging.WARNING)

# Args:
#     templates_dir_or_file (str): Path to the directory containing .docx files or path to a single docx document.


def check_fieldcode_delim(text):           # Goal: check if chevrons are not orphaned and not nested. This test should pass both on the run level and the paragraph level.
    chevrons = list(re.finditer(r'<<|>>', text))

    if not chevrons:  # No chevrons found, no error
        return False

    if chevrons[0].group() != "<<":  # Initial check for the first opening chevron
        logger.info(f"First chevron in para or run: Unexpected closing chevron at position {chevrons[0].start()}:::  '{text}'.")
        return f"First chevron in para or run: Unexpected closing chevron (>>) at position {chevrons[0].start()}:::  '{text}'."
    
    logger.setLevel(logging.WARNING)
    for i in range(1, len(chevrons), 2):  # Process pairs of chevrons
        open_chevron = chevrons[i - 1]
        close_chevron = chevrons[i]

        if open_chevron.group() != "<<" or close_chevron.group() != ">>":
            logger.info(f"Error at position {text[open_chevron.end():close_chevron.start()]}. Expected << >> pair.")
            return f"Error at position {text[open_chevron.end():close_chevron.start()]}. Expected << >> pair."

        # Process the chevron pair here...
        # For example, print the content between them:
        logger.info(f"Field code: {text[open_chevron.end():close_chevron.start()]}")

    # Check if there's an orphaned opening chevron
    if len(chevrons) % 2 != 0:
        return f"Orphaned opening chevron << at position {chevrons[-1].start()}: ::: '{text}'."
    
    return False


def realign_field_codes(docu: Document):

    # Task 2: Realign codes to the beginning of the same run. We can fix this without creating a new run by moving text to the end of the prev run, except the first run.

    for para in docu.paragraphs:
        #print (len(para.runs),  para.text[:70])
        if not para.runs:
            continue
        prevrun=para.runs[0]
            # print("First run length in the para: ", len(prevrun.text))
        for run in para.runs[1:]:       # Skip the first run. We cannot move anything to the left from there. todo: if one char length does [1:] throw error?
            #print("::"+ run.text +"::  len:", len (run.text))
            if "<<" in run.text and run.text[0:2] != "<<":   # Found a field code which DOES NOT align with the run. Realignment needed.

                thistextsplit = run.text.split("<<",1)
                logger.info("Moving:::::" + thistextsplit[0]+ ":::" + thistextsplit[1] + ":::::")
                prevrun.text += thistextsplit[0]
                run.text = "<<" + thistextsplit[1] 
                logger.info("After move:::::" + prevrun.text+ ":::" + run.text + ":::::")                
            prevrun=run


def preprocess_field_codes(docu: Document):

    # Goal: Fix field codes broken into two or more runs by moving all text that belongs to the same field code to a single run. 
    # At exit:
    #   if a run contains a field code, it will be at the start of the run. Except in the first run, due to the constraint that we don't want to create runs.
    #   field codes will not span runs 
    #   there may be more codes in each run.
    #   same runs exist as before. Unnecessary runs (i.e. runs that hosted only fraction of a code) may be cleared but still exist.

    # LLM based program can substitute field code with collected data only if Run-to-Code relationship is 1:many or 1:1, i.e. if no code is broken across many runs.
    # 1:1 provides for better formatting but works just as 1:many
    # This routine fixes if run:code is many:1. We do not have to create new runs for this.
    # It also attempts to improve formattability by moving the field code to the beginning of the run but does not create a new run. 


    # Task 1: Consolidate all runs that host the same field code into a single run, thereby fixing "many runs to the same field code" problem

    for para in docu.paragraphs: 

        i = 0

        while i < len(para.runs)-1:         # Testing the last run is not needed: it either contains no << or it also contains a >>. 

            if "<<" in para.runs[i].text and ">>" not in para.runs[i].text :    # skip this run if nothing to do i.e. this run has no opening mark or it contains the closing mark too
                codeopening = i   # Preserve the index of the run where the field code begins
                                    # This loop processes runs inside a field code so we don't have to bother with body text requirements.
                i += 1                                          # take the subsequent run
                innerrun_split = para.runs[i].text.split(">>",1)    # Slice the run into two pieces at the end of the code
                para.runs[codeopening].text += innerrun_split[0]    # Move the left slice (which is part of the field code) to the run that hosts the beginning of the same code. 
                                                                        # Combine the two runs if the code extends to one more run to the right
                while len(innerrun_split) <= 1:         # Check if the field code extends to the next run. 
                    para.runs[i].clear()                # if yes, clear all text as it has been copied to the run with the start mark
                    i += 1                                          # take the subsequent run
                    innerrun_split = para.runs[i].text.split(">>",1)    # Slice the run into two pieces at the end of the code
                    para.runs[codeopening].text += innerrun_split[0]    # Move the left slice (which is part of the field code) to the run that hosts the beginning of the same code. 
                                                                        # End of inner loop. Loop if not yet at the closing mark; exit if the field code finished.

                para.runs[codeopening].text += ">>"         # Split has chopped off the delimiter, add back.
                para.runs[i].text = innerrun_split[1]       # Remove text that has been copied from the rightmost run to the first run

            else:   
                i = i + 1       # take the next run


def create_fieldcode_docx(fieldcodes, contexts):
    """
    Create a docx file with a table based on the provided list of strings (fieldcodes).

    Args:
        fieldcodes (list): List of strings to include in the table.
        contexts (dict): Dictionary mapping fieldcodes to lists of context strings.

    Returns:
        Document object: The created docx.Document object.
    """
    def normalize_key(key):
        """Normalize a key by removing diacritics and replacing invalid characters."""
        normalized = unicodedata.normalize('NFKD', key).encode('ascii', 'ignore').decode('utf-8')
        normalized = ''.join(ch if ch in string.ascii_letters + string.digits + "_" else "_" for ch in normalized)
        return normalized

    def make_unique_key(base_key, used_keys):
        """Make a key unique by appending a number if necessary."""
        if base_key not in used_keys:
            return base_key
        counter = 1
        new_key = f"{base_key}_{counter}"
        while new_key in used_keys:
            counter += 1
            new_key = f"{base_key}_{counter}"
        return new_key

    # Initialize the document
    doc = Document()
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'

    # Create header row
    header_cells = table.rows[0].cells
    header_cells[0].text = "Serial"
    header_cells[1].text = "Key"
    header_cells[2].text = "Description"
    header_cells[3].text = "DocDisplay"
    header_cells[4].text = "UsedIn"

    used_keys = set()
    serial_counter = 1

    for fieldcode in fieldcodes:
        # Check if the fieldcode is a valid Python identifier

        if fieldcode.isidentifier() and all(ch in string.ascii_letters + string.digits + "_" for ch in fieldcode):
            key = fieldcode
            # print(f"valid identifyer {fieldcode}")
        else:
            normalized_key = normalize_key(fieldcode)
            key = make_unique_key(normalized_key, used_keys)

        # Track key usage
        used_keys.add(key)

        # Add a new row
        row_cells = table.add_row().cells
        row_cells[0].text = str(serial_counter)
        row_cells[2].text = ""  # Description column is empty
        row_cells[3].text = fieldcode
        if len(contexts[fieldcode]) == 1:
            row_cells[4].text = contexts[fieldcode][0]
        else:       # Add all contexts to the "UsedIn" column
            for ctx in contexts[fieldcode]:
                run = row_cells[4].paragraphs[0].add_run()
                run.text = ctx
                run_sep = row_cells[4].paragraphs[0].add_run()
                run_sep.text = ";  "
                run.font.color.rgb = RGBColor(255, 0, 0)  # Set the font color to red
                #  run.font.color.rgb = RGBColor(0, 0, 0)  # Set the font color to red

        # Format the key cell if it was normalized
        if key != fieldcode:
            run = row_cells[1].paragraphs[0].add_run()
            run.text = key
            run.font.color.rgb = RGBColor(255, 0, 0)  # Set the font color to red
            run.bold = True
        else:
            row_cells[1].paragraphs[0].add_run(key)

        serial_counter += 1

    return doc if serial_counter > 1 else None

def extract_fieldcodes_and_contexts(doc):
    """
    Extract unique field codes and their contexts from a Word document.

    Args:
        doc (Document): A docx.Document object to extract field codes from.

    Returns:
        list: A list of unique field codes.
        dict: A dictionary mapping each field code to a list of its contexts.
    """
    fieldcodes = []
    contexts = defaultdict(list)
    text = "".join([p.text for p in doc.paragraphs])

    for match in re.finditer(r"<<([^<>]+)>>", text):
        fieldcode = match.group(1)
        start, end = match.start(), match.end()
        context = text[max(0, start - 15):min(len(text), end + 15)]

        if fieldcode not in fieldcodes:
            fieldcodes.append(fieldcode)
        contexts[fieldcode].append(context)

    return fieldcodes, contexts




def do_checks_and_fixes(input_filename, input_dir, extract_fieldcodes = False):

    
    # Create the PreProc subdirectory path
    preproc_dir = os.path.join(input_dir, "PreProc")
    os.makedirs(preproc_dir, exist_ok=True)
    
    # Modify the filename to include "proc" before the extension
    base_name, ext = os.path.splitext(input_filename)
    output_filename = f"{base_name}_proc{ext}"

    doc_output_fullpath = os.path.join(preproc_dir, output_filename)

    # doc_output_fullpath = os.path.join(templates_dir, file_name + ".proc.docx") # dirname plus output plus filename 
    docuT = Document(os.path.join(input_dir,input_filename))
    docuT.core_properties.category = "Preprocessed template for AI chatbot Lisa"
    docuT.save(doc_output_fullpath)

    print("  Testing field code delimiters for each paragraph. ")

    checks= []
    for para in docuT.paragraphs:
        checks.append(check_fieldcode_delim(para.text))
    if any(checks):
        print(f"{input_filename} cannot be processed. Please ensure that opening and closing marks of the same field code are in the same paragraph. ") 
        return checks
    print("  Tested if each paragraph includes pairs of opening and closing marks. Result: PASSED")

    #print ("Checking on the run level")
    checks= []
    for para in docuT.paragraphs:
        #logger.info("Starting a new para '%s'", para.text[0:40])
        for run in para.runs:
            checks.append(check_fieldcode_delim(run.text))
            # print(f"Latest checks[-1]: {checks[-1]}")
    print("  Tested if each subparagraph (run) includes pairs of opening and closing marks. " + ("FAILED but will be fixed if possible." if any(checks) else "PASSED"))


    ###### Step 1: combine runs into the same run if they host the same field code.
    print("  Start preprocessing Step 1 - copy all parts of the same field code into the same run")
    preprocess_field_codes(docuT)

    ###### Step 2: realign field codes to the beginning of a run if possible
    print("  Start preprocessing Step 2 - realignment of field codes to the beginning of runs")
    realign_field_codes(docuT)

    docuT.save(doc_output_fullpath)

    #print("  Preprocessed and saved to ", output_filename)

    # print ("  After preprocessing: Checking validity of field codes on the run (subparagraph) level")
    checks= []
    for para in docuT.paragraphs:
        #    logger.info("Starting a new para '%s'", para.text[0:40])
        lastchar_in_prev = ""
        for run in para.runs:
            checks.append(check_fieldcode_delim(run.text))
            if checks[-1] and lastchar_in_prev == "<" and run.text and run.text[0] == "<" and not run.text.startswith("<<") :   # if a marker is separated into two different runs (i.e. the run starts with a char that completes the previous run's last char to make up a field code mark)
                print("  Likely error:  ", run.text)
                checks[-1] += " --- Two consecutive '<' signs were found. If they were meant to be a field code opening mark then it was missed as it was split into two different runs. Check accidental formatting."       #   signal suspicion but do not try to fix. Likely but not surely a formatting error. 
            if checks[-1] and lastchar_in_prev == ">" and run.text and run.text[0] == ">" and not run.text.startswith(">>") :   # if a marker is separated into two different runs (i.e. the run starts with a char that completes the previous run's last char to make up a field code mark)
                checks[-1] += " --- Two consecutive '>' signs were found. If they were meant to be a field code closing mark then it was missed as it was split into two different runs. Check accidental formatting."       #   signal suspicion but do not try to fix. Likely but not surely a formatting error. 
            if run.text:
                lastchar_in_prev = run.text[-1]     # Preserve last char of this run if this is not an empty run. Used to test opening marks separated into two runs and closing marks separated into two runs accidentally 

    #print(checks)
    if not any(checks):
        print("  Test PASSED." +  f"  Saved {doc_output_fullpath}.")

        if extract_fieldcodes:
            fieldcode_filename = f"{base_name}_flds{ext}"
            doc_fields_output_fullpath = os.path.join(preproc_dir, fieldcode_filename)

            fieldcodes, contexts = extract_fieldcodes_and_contexts(docuT)
            doc_fld = create_fieldcode_docx(fieldcodes, contexts)
            if doc_fld:
                doc_fld.core_properties.category = "Field codes from template for AI chatbot Lisa"
                doc_fld.save(doc_fields_output_fullpath)
                print(f"  Saved field list in {doc_fields_output_fullpath}.")

    else:
        print("  Test FAILED. Field code marks and their formatting should be checked." +  f"  Saved {doc_output_fullpath}.")

    return checks


def main():

    def printusage():
        print("Usage: python templatepreproc.py [-f] <path>")
        print("       <path> points to either a single template or a directory with templates.")
        print("       -f: save the list of field codes found in the template. Output: [base_name]_flds.docx")
        print("Templates are docx files with field codes between << and >>")
        print("Output is saved in the PreProc subdirectory. If file exists it is overwritten.")
        sys.exit(1)

    # Ensure a command-line argument is provided
    if len(sys.argv) < 2:
        printusage()

    # Get the directory path from the command-line argument

    extract_fieldcodes = (sys.argv[1] == "-f")

    if len(sys.argv) < 3 and extract_fieldcodes:
        printusage()

    input_path = sys.argv[2] if extract_fieldcodes else sys.argv[1]

    input_abspath = os.path.abspath(input_path)

    # Check if the provided path is a valid directory
    if os.path.isdir(input_path):
        print(f"Processing all .docx files in directory '{input_path}'.")

        #   Iterate over all .docx files in the given directory and call a function for each.
        for file_name in os.listdir(input_abspath):
            if file_name.endswith(".docx"):
                print()
                print(f"########  Processing .docx file '{file_name}':")            
                logger.info("Processing template '%s' ",  file_name )

                errors = do_checks_and_fixes(file_name, input_abspath, extract_fieldcodes) # input_fullpath = os.path.join(input_abspath, file_name))
                if any(errors):
                    print("This is a list of errors in the input template with a short context to help locate issues:")
                    print([item for item in errors if item])        # print only if error is not False
                print()

        print(f"Completed all docx documents in {input_abspath}.")

    elif os.path.exists(input_path) and input_path.endswith(".docx"):
        print()
        print(f"#########  Processing .docx file '{input_path}':")
        logger.info("Processing template '%s' ",  input_abspath )

        errors = do_checks_and_fixes( os.path.basename(input_abspath), os.path.dirname(input_abspath), extract_fieldcodes )
        if any(errors):
            print("This is a list of errors in the input template with a short context to help locate issues:")
            print([item for item in errors if item])        # print only if the check item is not False
        else: 
            print(f"Completed template {input_path}")


    else:
        print("Error: path is not a directory or a docx file.")
        sys.exit(1)

if __name__ == "__main__":
    main()